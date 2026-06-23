import io
from types import SimpleNamespace

import fitz
import pytest
from docx import Document

import app.ingestion.vision as vision_mod
from app.ingestion.detector import get_extractor
from app.ingestion.docx import extract_docx
from app.ingestion.models import ElementType


def _png_bytes() -> bytes:
    pix = fitz.Pixmap(fitz.csRGB, fitz.IRect(0, 0, 64, 64))
    pix.clear_with(200)
    return pix.tobytes("png")


def _make_docx() -> bytes:
    doc = Document()
    doc.add_heading("Introduction", level=1)
    doc.add_paragraph("This is the intro paragraph.")
    doc.add_heading("Methods", level=1)
    doc.add_paragraph("We used method X.")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Alpha"
    table.cell(1, 1).text = "42"
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _make_docx_with_image() -> bytes:
    doc = Document()
    doc.add_heading("Diagrams", level=1)
    doc.add_picture(io.BytesIO(_png_bytes()))
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_extract_docx_sections_and_table() -> None:
    result = extract_docx("doc.docx", _make_docx())
    assert result.format == "docx"
    types = [e.element_type for e in result.elements]
    assert ElementType.HEADING in types
    assert ElementType.TABLE in types
    methods_paras = [
        e
        for e in result.elements
        if e.element_type == ElementType.PARAGRAPH and e.location.section == "Methods"
    ]
    assert any("method X" in e.text for e in methods_paras)
    table_el = next(e for e in result.elements if e.element_type == ElementType.TABLE)
    assert "Name" in table_el.text and "42" in table_el.text
    assert table_el.location.section == "Methods"


def test_extract_docx_block_index_anchor() -> None:
    result = extract_docx("doc.docx", _make_docx())
    assert all(e.location.block_index is not None for e in result.elements)
    indices = [e.location.block_index for e in result.elements]
    assert indices == sorted(indices)  # type: ignore


def test_extract_docx_image_via_vision(monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setattr(vision_mod, "describe_image", lambda blob, mime: "an architecture diagram")
    monkeypatch.setattr(
        vision_mod,
        "get_settings",
        lambda: SimpleNamespace(vision_min_image_kb=0, vision_max_images=10),
    )
    result = extract_docx("doc.docx", _make_docx_with_image())
    image_els = [e for e in result.elements if e.element_type == ElementType.IMAGE]
    assert any("an architecture diagram" in e.text for e in image_els)
    assert image_els[0].location.section == "Diagrams"


def test_docx_extractor_is_registered() -> None:
    assert get_extractor("docx") is extract_docx
